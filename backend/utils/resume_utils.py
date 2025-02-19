import pdfplumber
import docx
import spacy
from sentence_transformers import SentenceTransformer, util
from io import BytesIO

nlp = spacy.load("en_core_web_sm")
bert_model = SentenceTransformer('all-MiniLM-L6-v2')

async def extract_text_from_pdf(file):
    text = ""
    file_data = await file.read()
    with pdfplumber.open(BytesIO(file_data)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text.strip() if text.strip() else "No readable text found in PDF."

async def extract_text_from_docx(file):
    file_data = await file.read()
    doc = docx.Document(BytesIO(file_data))
    return "\n".join([para.text for para in doc.paragraphs]) if doc.paragraphs else "No text found in DOCX file."

def extract_keywords(text):
    doc = nlp(text)
    return [token.lemma_ for token in doc if token.pos_ in ["NOUN", "PROPN"]]

def calculate_similarity(resume_text, job_desc):
    resume_embedding = bert_model.encode(resume_text, convert_to_tensor=True)
    job_embedding = bert_model.encode(job_desc, convert_to_tensor=True)
    similarity_score = util.pytorch_cos_sim(resume_embedding, job_embedding).item()
    return round(similarity_score * 100, 2)

async def analyze_resume(file, job_description):
    file_extension = file.filename.split(".")[-1]
    
    if file_extension == "pdf":
        resume_text = await extract_text_from_pdf(file)
    elif file_extension == "docx":
        resume_text = await extract_text_from_docx(file)
    else:
        return {"error": "Unsupported file format. Upload a PDF or DOCX file."}

    keywords = extract_keywords(resume_text)
    similarity = calculate_similarity(resume_text, job_description)

    return {
        "resume_keywords": keywords,
        "match_score": similarity,
        "suggestion": "Add missing skills from job description to improve your match."
    }
